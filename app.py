from flask import Flask, request, jsonify, render_template, Response
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import OllamaLLM
from langchain_core.documents import Document
import os
import uuid
from werkzeug.utils import secure_filename
import PyPDF2
import docx

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = "uploads"
app.config["ALLOWED_EXTENSIONS"] = {"txt", "pdf", "docx", "doc"}

# Create uploads directory if it doesn't exist
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

# Store uploaded documents and their embeddings
documents_store = {}

# Initialize the chatbot components
context = """
This is a generic chatbot. It can answer questions based on the information provided to it.
For now, it knows a little about AI: Artificial Intelligence (AI) is the simulation of human intelligence in machines.
AI systems can perform tasks like learning, problem-solving, and decision-making.
"""
documents = [Document(page_content=context)]

text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
docs = text_splitter.split_documents(documents)
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
vector_store = FAISS.from_documents(docs, embeddings)
llm = OllamaLLM(model="deepseek-r1:7b", base_url="http://localhost:11434")


def allowed_file(filename):
    return (
        "." in filename
        and filename.rsplit(".", 1)[1].lower() in app.config["ALLOWED_EXTENSIONS"]
    )


def extract_text_from_file(file_path):
    """Extract text from different file types"""
    file_extension = file_path.split(".")[-1].lower()

    if file_extension == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    elif file_extension == "pdf":
        text = ""
        with open(file_path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num in range(len(pdf_reader.pages)):
                text += pdf_reader.pages[page_num].extract_text()
        return text
    elif file_extension in ["docx", "doc"]:
        doc = docx.Document(file_path)
        return " ".join([paragraph.text for paragraph in doc.paragraphs])
    return ""


def process_document(file_path, filename):
    """Process document and create embeddings"""
    doc_text = extract_text_from_file(file_path)
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = text_splitter.split_text(doc_text)
    doc_chunks = [Document(page_content=chunk) for chunk in chunks]
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )
    vectorstore = FAISS.from_documents(doc_chunks, embeddings)
    doc_id = str(uuid.uuid4())
    documents_store[doc_id] = {
        "id": doc_id,
        "filename": filename,
        "vectorstore": vectorstore,
        "raw_text": doc_text,
    }
    return doc_id


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/upload-documents", methods=["POST"])
def upload_documents():
    if "documents" not in request.files:
        return jsonify({"error": "No documents provided"}), 400

    files = request.files.getlist("documents")
    uploaded_docs = []

    for file in files:
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            file.save(file_path)
            doc_id = process_document(file_path, filename)
            uploaded_docs.append({"id": doc_id, "filename": filename})

    return jsonify({"documents": uploaded_docs})


@app.route("/delete-document/<doc_id>", methods=["DELETE"])
def delete_document(doc_id):
    if doc_id in documents_store:
        del documents_store[doc_id]
        return jsonify({"success": True})
    return jsonify({"error": "Document not found"}), 404


@app.route("/chat", methods=["POST"])
def chat():
    print("Chat request received")
    user_input = request.json.get("message", "")
    document_ids = request.json.get("documentIds", [])
    stream_mode = request.json.get("stream", False)
    print(f"User input: {user_input}, Stream mode: {stream_mode}")

    try:
        if user_input.strip():
            if document_ids:
                context = ""
                for doc_id in document_ids:
                    if doc_id in documents_store:
                        doc = documents_store[doc_id]
                        vectorstore = doc["vectorstore"]
                        relevant_docs = vectorstore.similarity_search(user_input, k=3)
                        context += f"\n\n--- From {doc['filename']} ---\n"
                        for doc_chunk in relevant_docs:
                            context += doc_chunk.page_content + "\n"
                if not context:
                    context = "No relevant information found in the uploaded documents."
            else:
                retriever = vector_store.as_retriever(search_kwargs={"k": 3})
                retrieved_docs = retriever.invoke(user_input)
                context = "\n".join([doc.page_content for doc in retrieved_docs])

            prompt = f"Based on this context:\n{context}\n\nAnswer the question: {user_input}"

            if stream_mode:
                print("Using streaming mode for response")

                def generate():
                    try:
                        for chunk in llm.stream(prompt):
                            yield f"data: {chunk}\n\n"
                        yield "data: [DONE]\n\n"
                    except Exception as e:
                        print(f"Error in stream generation: {str(e)}")
                        yield f"data: Error in stream: {str(e)}\n\n"
                        yield "data: [DONE]\n\n"

                return Response(generate(), mimetype="text/event-stream")
            else:
                print("Using non-streaming mode for response")
                chat_output = llm.invoke(prompt)
                return jsonify({"response": chat_output})
        else:
            return jsonify({"response": "I'm not sure how to respond to that."})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"response": f"Error processing your request: {str(e)}"})


if __name__ == "__main__":
    print(f"Loaded and split {len(docs)} document chunks.")
    print("Starting chatbot server...")
    app.run(debug=True, host="0.0.0.0", port=5000)
